
import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def fix_quotes(text):
    result = ''
    i = 0
    open_quote = True
    while i < len(text):
        if text[i] == '"':
            if open_quote:
                result += '\u201c'  # 中文左引号
                open_quote = False
            else:
                result += '\u201d'  # 中文右引号
                open_quote = True
            i += 1
        elif text[i] == "'":
            if open_quote:
                result += '\u2018'
                open_quote = False
            else:
                result += '\u2019'
                open_quote = True
            i += 1
        else:
            result += text[i]
            if text[i].isspace() or i == len(text) - 1:
                open_quote = True
            i += 1
    return result

def split_text(text):
    text = fix_quotes(text)
    parts = []
    current = ''
    is_ascii = None
    
    for char in text:
        char_is_ascii = bool(re.match(r'[0-9a-zA-Z\.\-\:\/\,\%\s\+\(\)]+', char))
        if is_ascii is None:
            is_ascii = char_is_ascii
            current = char
        elif is_ascii == char_is_ascii:
            current += char
        else:
            parts.append((current, is_ascii))
            current = char
            is_ascii = char_is_ascii
    
    if current:
        parts.append((current, is_ascii))
    return parts

def format_official_document(input_path, output_path, doc_type='general'):
    doc = Document(input_path)
    
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            para.paragraph_format.line_spacing = Pt(28)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue
        
        para.style = doc.styles['Normal']
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(28)
        
        original_text = para.text
        for run in para.runs:
            run.text = ''
        
        is_title = False
        chinese_font = '仿宋_GB2312'
        text_stripped = text.strip()
        
        if doc_type == '附件' and re.match(r'^附件\d+$', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            chinese_font = '黑体'
            is_title = True
            para.add_run(original_text).bold = True
        
        elif (any(k in text_stripped for k in ['请示', '报告', '通知', '方案', '总结', '规定', '办法', '意见', '函', '纪要', '明细表']) 
              and len(text_stripped) < 50 and len(text_stripped) > 10 
              and '，' not in text_stripped and '。' not in text_stripped
              and not re.match(r'^[一二三四五六七八九十]+、', text_stripped) 
              and not re.match(r'^（[一二三四五六七八九十]+）', text_stripped) 
              and not re.match(r'^\d+\.', text_stripped)):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            chinese_font = '方正小标宋简体'
            is_title = True
            run = para.add_run(original_text)
            run.font.name = '方正小标宋简体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            run.font.size = Pt(22)
            run.bold = False
            continue
        
        elif text_stripped.endswith('：') and any(k in text_stripped for k in ['领导', '公司', '局', '委', '办', '同志']):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif re.match(r'^[一二三四五六七八九十]+、', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '黑体'
            is_title = True
            run = para.add_run(original_text)
            run.bold = False  # 明确设置为不加粗
        
        elif re.match(r'^（[一二三四五六七八九十]+）', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '楷体_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif re.match(r'^\d+\.\s*\S', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif re.match(r'^（\d+）', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif text_stripped.startswith('附件：') or text_stripped.startswith('附：'):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0.74)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif (re.match(r'^[2-9]\.', text_stripped) or re.match(r'^\d+\.\s*\S', text_stripped)) and len(text_stripped) < 30:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0.74)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif '妥否' in text_stripped and '请批示' in text_stripped:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0.74)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        elif doc_type == '请示' or doc_type == '函':
            is_signature = (len(text_stripped) < 30 and 
                          any(s in text_stripped for s in ['部', '司', '局', '委', '办', '公司', '集团', '所', '院', '中心', '组', '室']))
            is_date = re.match(r'^\d{4}年\d{1,2}月\d{1,2}日$', text_stripped) or \
                     re.match(r'^[一二三四五六七八九十]{4}年[一二三四五六七八九十]{1,2}月[一二三四五六七八九十]{1,2}日$', text_stripped)
            
            if is_signature or is_date:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                para.paragraph_format.first_line_indent = Cm(0)
                chinese_font = '仿宋_GB2312'
                is_title = True
                para.add_run(original_text)
        
        if not is_title:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            
            is_date = re.match(r'^\d{4}年\d{1,2}月\d{1,2}日$', text_stripped) or \
                     re.match(r'^[一二三四五六七八九十]{4}年[一二三四五六七八九十]{1,2}月[一二三四五六七八九十]{1,2}日$', text_stripped)
            if text_stripped.endswith(')') or text_stripped.endswith('）'):
                pass
            elif text_stripped and not text_stripped.endswith(('。', '！', '？', '：', '；', '，', '日')) and len(text_stripped) > 3 and not is_date:
                original_text += '。'
            
            para.add_run(original_text)
        
        original_text = fix_quotes(original_text)
        for run in para.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = False  # 先清除所有加粗
            
        for run in para.runs:
            run.text = ''
        
        parts = split_text(original_text)
        if len(parts) == 1:
            text_part, is_ascii = parts[0]
            if is_ascii:
                run = para.add_run(text_part)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = False
            else:
                run = para.add_run(text_part)
                run.font.name = chinese_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = False
        else:
            for text_part, is_ascii in parts:
                new_run = para.add_run(text_part)
                new_run.font.size = Pt(16)
                new_run.font.color.rgb = RGBColor(0, 0, 0)
                new_run.bold = False
                if is_ascii:
                    new_run.font.name = 'Times New Roman'
                    new_run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                else:
                    new_run.font.name = chinese_font
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
    
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    b = OxmlElement(f'w:{bn}')
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), '4')
                    b.set(qn('w:color'), '000000')
                    tcBorders.append(b)
                tcPr.append(tcBorders)
                
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = Pt(28)
                    for r in p.runs:
                        r.font.name = '仿宋_GB2312'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        r.font.size = Pt(16)
    
    doc.save(output_path)
    print(f"✅ 文档格式化完成：{output_path}")
    print("✅ 格式规范：")
    print("   • 一级标题：黑体三号，不加粗，缩进1.27cm")
    print("   • 三级标题：仿宋_GB2312三号，缩进1.27cm")
    print("   • 数字/英文：Times New Roman三号")
    print("   • 日期后面不加句号")
    print("   • 保留原文档空行结构，只格式化不改变段落分隔")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("用法：python format_docx.py 输入文件.docx 输出文件.docx 文档类型")
        sys.exit(1)
    
    try:
        format_official_document(sys.argv[1], sys.argv[2], sys.argv[3])
    except Exception as e:
        print(f"❌ 格式化失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
