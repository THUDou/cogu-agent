
import os
import tempfile

COMMAND = "pdf_to_word"
DESCRIPTION = "Convert PDF to Word document (.docx) with image enhancement"
CATEGORY = "organize"

PARAMS = [
    {"name": "input",  "type": "str", "required": True,  "help": "Input PDF path"},
    {"name": "output", "type": "str", "required": True,  "help": "Output Word path (.docx)"},
    {"name": "pages",  "type": "json", "required": False,
     "help": "Page range [start, end] (0-based, default all)"},
    {"name": "method", "type": "str", "required": False, "default": "auto",
     "choices": ["auto", "pdf2docx", "libreoffice"],
     "help": "Conversion method"},
    {"name": "enhance_images", "type": "bool", "required": False, "default": True,
     "help": "Enable image enhancement (extract and re-insert missing images)"},
]



def _extract_pdf_images(input_path, page_range=None, min_size=50):
    import fitz

    doc = fitz.open(input_path)
    total_pages = len(doc)

    if page_range and len(page_range) >= 2:
        start, end = page_range[0], min(page_range[1], total_pages - 1)
        page_indices = range(start, end + 1)
    else:
        page_indices = range(total_pages)

    images = []
    seen_xrefs = set()

    for page_idx in page_indices:
        page = doc[page_idx]
        page_width = page.rect.width
        page_height = page.rect.height

        img_list = page.get_images(full=True)

        for img_info in img_list:
            xref = img_info[0]
            if xref in seen_xrefs:
                continue

            img_width = img_info[2]
            img_height = img_info[3]

            if img_width < min_size and img_height < min_size:
                continue

            bbox = None
            try:
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    rect = img_rects[0]
                    bbox = {
                        "x0": rect.x0,
                        "y0": rect.y0,
                        "x1": rect.x1,
                        "y1": rect.y1,
                        "rel_y": rect.y0 / page_height if page_height > 0 else 0,
                    }
            except Exception:
                pass

            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:  # CMYK / CMYK+Alpha → RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)

                image_bytes = pix.tobytes("png")
                seen_xrefs.add(xref)

                images.append({
                    "page": page_idx,
                    "xref": xref,
                    "width": pix.width,
                    "height": pix.height,
                    "bbox": bbox,
                    "image_bytes": image_bytes,
                    "ext": "png",
                    "page_width": page_width,
                    "page_height": page_height,
                })
                pix = None
            except Exception:
                seen_xrefs.add(xref)
                continue

    doc.close()
    return images



def _count_docx_images(docx_path):
    try:
        from docx import Document
        doc = Document(docx_path)
        count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                count += 1
        return count
    except Exception:
        return -1


def _insert_missing_images(docx_path, pdf_images, output_path):
    from docx import Document
    from docx.shared import Inches, Pt
    import io

    doc = Document(docx_path)

    existing_images = set()
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                existing_images.add(len(img_data))
            except Exception:
                pass

    docx_image_count = len(existing_images)

    missing_images = []
    for img in pdf_images:
        img_size = len(img["image_bytes"])
        found = False
        for existing_size in existing_images:
            ratio = img_size / existing_size if existing_size > 0 else 0
            if 0.5 < ratio < 2.0:
                found = True
                break
        if not found:
            missing_images.append(img)

    if not missing_images:
        return {
            "inserted": 0,
            "total_pdf": len(pdf_images),
            "total_docx": docx_image_count,
            "message": "Word 中图片完整，无需补全",
        }

    from collections import defaultdict
    page_images = defaultdict(list)
    for img in missing_images:
        page_images[img["page"]].append(img)

    total_paragraphs = len(doc.paragraphs)
    total_pdf_pages = max(img["page"] for img in pdf_images) + 1 if pdf_images else 1
    paragraphs_per_page = max(1, total_paragraphs // total_pdf_pages)

    inserted_count = 0

    for page_idx in sorted(page_images.keys(), reverse=True):
        imgs = sorted(page_images[page_idx],
                       key=lambda x: x["bbox"]["rel_y"] if x.get("bbox") else 0.5,
                       reverse=True)

        for img in imgs:
            rel_y = img["bbox"]["rel_y"] if img.get("bbox") else 0.5
            base_para_idx = int(page_idx * paragraphs_per_page)
            insert_idx = min(
                base_para_idx + int(rel_y * paragraphs_per_page),
                total_paragraphs
            )

            if img.get("bbox") and img.get("page_width"):
                img_width_ratio = (img["bbox"]["x1"] - img["bbox"]["x0"]) / img["page_width"]
                display_width = min(6.0, max(1.0, img_width_ratio * 6.5))
            else:
                display_width = min(6.0, img["width"] / 96.0)  # 96 DPI
                display_width = max(1.0, display_width)

            try:
                if insert_idx < len(doc.paragraphs):
                    target_para = doc.paragraphs[insert_idx]
                    new_para = target_para.insert_paragraph_before("")
                else:
                    new_para = doc.add_paragraph("")

                run = new_para.add_run()
                run.add_picture(io.BytesIO(img["image_bytes"]), width=Inches(display_width))
                inserted_count += 1
            except Exception:
                continue

    doc.save(output_path)

    return {
        "inserted": inserted_count,
        "total_pdf": len(pdf_images),
        "total_docx": docx_image_count,
        "missing_detected": len(missing_images),
    }



def _convert_pdf2docx(input_path, output_path, pages):
    from pdf2docx import Converter

    cv = Converter(input_path)
    if pages and len(pages) >= 2:
        start, end = pages[0], pages[1]
        cv.convert(output_path, start=start, end=end + 1)
        page_count = end - start + 1
    else:
        cv.convert(output_path)
        page_count = len(cv.pages) if hasattr(cv, 'pages') else -1
    cv.close()
    return page_count


def _convert_libreoffice(input_path, output_path):
    import shutil
    import subprocess

    lo_path = shutil.which("soffice") or shutil.which("libreoffice")
    if not lo_path:
        raise FileNotFoundError("LibreOffice 未安装")

    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            lo_path, "--headless",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            input_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")

        base_name = os.path.splitext(os.path.basename(input_path))[0]
        tmp_output = os.path.join(tmpdir, base_name + ".docx")
        if not os.path.exists(tmp_output):
            raise FileNotFoundError(f"LibreOffice 输出文件未找到: {tmp_output}")

        shutil.move(tmp_output, output_path)

    import fitz
    doc = fitz.open(input_path)
    page_count = len(doc)
    doc.close()
    return page_count



def handler(params):
    input_path = params["input"]
    output_path = params["output"]
    pages = params.get("pages", None)
    method = params.get("method", "auto")
    enhance = params.get("enhance_images", True)

    if not os.path.exists(input_path):
        raise FileNotFoundError(f"文件不存在: {input_path}")

    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    pdf_images = []
    if enhance:
        try:
            pdf_images = _extract_pdf_images(input_path, page_range=pages)
        except Exception:
            pdf_images = []

    engine_used = "none"
    page_count = 0

    if method in ("auto", "pdf2docx"):
        try:
            page_count = _convert_pdf2docx(input_path, output_path, pages)
            engine_used = "pdf2docx"
        except ImportError:
            if method == "pdf2docx":
                raise ImportError("pdf2docx 未安装，请运行: pip install pdf2docx")
        except Exception:
            if method == "pdf2docx":
                raise

    if engine_used == "none" and method in ("auto", "libreoffice"):
        try:
            page_count = _convert_libreoffice(input_path, output_path)
            engine_used = "libreoffice"
        except FileNotFoundError:
            if method == "libreoffice":
                raise FileNotFoundError("LibreOffice 未安装")
        except Exception:
            if method == "libreoffice":
                raise

    if engine_used == "none":
        raise RuntimeError("无可用的 PDF 转 Word 引擎。请安装 pdf2docx 或 LibreOffice。")

    image_enhance_result = None
    if enhance and pdf_images and os.path.exists(output_path):
        try:
            docx_img_count = _count_docx_images(output_path)
            pdf_img_count = len(pdf_images)

            if docx_img_count >= 0 and pdf_img_count > docx_img_count:
                image_enhance_result = _insert_missing_images(
                    output_path, pdf_images, output_path
                )
            else:
                image_enhance_result = {
                    "inserted": 0,
                    "total_pdf": pdf_img_count,
                    "total_docx": docx_img_count,
                    "message": "图片完整，无需补全",
                }
        except Exception as e:
            image_enhance_result = {
                "inserted": 0,
                "error": str(e),
            }

    result = {
        "success": True,
        "output": output_path,
        "pages_converted": page_count,
        "engine": engine_used,
        "file_size": os.path.getsize(output_path),
    }

    if image_enhance_result is not None:
        result["image_enhancement"] = image_enhance_result

    return result


if __name__ == "__main__":
    from pdfkit.base import main
    main(handler, params_schema=PARAMS, description=DESCRIPTION)
